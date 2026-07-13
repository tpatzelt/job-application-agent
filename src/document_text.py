from __future__ import annotations

import io
import logging

logger = logging.getLogger(__name__)

# Reject extracted texts shorter than this; a CV or letter this short is
# almost certainly a failed extraction (scanned PDF, wrong file, ...).
MIN_DOCUMENT_CHARS = 40


class DocumentExtractionError(RuntimeError):
    """Raised when no usable text could be extracted from an upload."""


def extract_text(data: bytes, file_name: str, mime_type: str = "") -> str:
    """Extract plain text from an uploaded document (PDF, DOCX, or text)."""
    name = file_name.lower()
    if name.endswith(".pdf") or "pdf" in mime_type:
        text = _extract_pdf(data)
    elif name.endswith(".docx") or "officedocument.wordprocessingml" in mime_type:
        text = _extract_docx(data)
    elif name.endswith(".doc"):
        raise DocumentExtractionError(
            "Legacy .doc files are not supported; please export as PDF or DOCX."
        )
    else:
        text = data.decode("utf-8", errors="replace")
    text = _normalize(text)
    if len(text) < MIN_DOCUMENT_CHARS:
        raise DocumentExtractionError(
            "Could not extract readable text from the document. If it is a "
            "scanned PDF, please send a text-based PDF, DOCX, or plain text."
        )
    return text


def _extract_pdf(data: bytes) -> str:
    try:
        from pypdf import PdfReader
    except ImportError as exc:  # pragma: no cover - dependency is declared
        raise DocumentExtractionError(f"PDF support unavailable: {exc}") from exc
    try:
        reader = PdfReader(io.BytesIO(data))
        return "\n".join(page.extract_text() or "" for page in reader.pages)
    except Exception as exc:
        logger.warning("PDF extraction failed: %s", exc)
        raise DocumentExtractionError(f"Could not read PDF: {exc}") from exc


def _extract_docx(data: bytes) -> str:
    try:
        import docx
    except ImportError as exc:  # pragma: no cover - dependency is declared
        raise DocumentExtractionError(f"DOCX support unavailable: {exc}") from exc
    try:
        document = docx.Document(io.BytesIO(data))
        parts = [paragraph.text for paragraph in document.paragraphs]
        for table in document.tables:
            for row in table.rows:
                parts.extend(cell.text for cell in row.cells)
        return "\n".join(parts)
    except Exception as exc:
        logger.warning("DOCX extraction failed: %s", exc)
        raise DocumentExtractionError(f"Could not read DOCX: {exc}") from exc


def _normalize(text: str) -> str:
    lines = [line.strip() for line in text.splitlines()]
    return "\n".join(line for line in lines if line).strip()
